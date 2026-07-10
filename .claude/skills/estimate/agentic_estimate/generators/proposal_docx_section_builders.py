"""Helper functions for building proposal DOCX sections."""

from .estimate_render_helpers import get_roles, task_role_md
from .proposal_docx_helpers import add_styled_table, flatten_tasks, sum_option_md


def _add_centered(doc, text: str, size: int, bold: bool = False, space_after: int = 12):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(space_after)


def add_cover_page(doc, data: dict, config: dict):
    _add_centered(doc, data.get("project_name", "Project Proposal"), 28, bold=True)
    _add_centered(doc, "Project Proposal", 16, space_after=24)
    date_text = data.get("generated_date") or config.get("date", "")
    _add_centered(doc, f"Date: {date_text}", 12)
    doc.add_page_break()


def add_toc(doc):
    from docx.shared import Pt, RGBColor

    heading = doc.add_heading("Table of Contents", level=1)
    heading.space_after = Pt(12)
    toc_note = doc.add_paragraph("[Right-click → Update Field to generate TOC]")
    toc_note.runs[0].font.italic = True
    toc_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_page_break()


def add_executive_summary(doc, data: dict):
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(data.get("description", "Project estimation and scope overview"))
    doc.add_paragraph()
    if not data.get("options"):
        return
    total_md = sum_option_md(data["options"][0])
    active_roles, _ = get_roles(data)
    team_size = len(active_roles)
    duration = round(total_md / max(team_size, 1) / 5) if total_md > 0 else 0
    params = data.get("parameters", {})
    cost = int(total_md * params.get("cost_per_md", 0))
    currency = params.get("currency", "USD")
    conf = data.get("confidence", {})
    rows = [
        ["Total Effort", f"{total_md:.1f} man-days"],
        ["Team Size", f"{team_size} roles"],
        ["Estimated Duration", f"{duration} weeks (assuming 5 MD/week)"],
        ["Total Cost", f"{cost:,} {currency}"],
        ["Confidence", f"{conf.get('level', 'N/A')} (±{conf.get('range_pct', 0)}%)"],
    ]
    add_styled_table(doc, ["Metric", "Value"], rows)
    doc.add_paragraph()


def add_scope(doc, data: dict):
    doc.add_heading("2. Scope of Work", level=1)
    doc.add_heading("In Scope", level=2)
    if data.get("options"):
        option = data["options"][0]
        categories = option.get("categories", [])
        if categories:
            for cat in categories:
                doc.add_paragraph(f"{cat.get('name', 'Unnamed')}", style="List Bullet")
                for task in cat.get("tasks", []):
                    doc.add_paragraph(f"  • {task.get('name', '')}", style="List Bullet 2")
        else:
            for task in option.get("tasks", []):
                doc.add_paragraph(task.get("name", ""), style="List Bullet")
    doc.add_paragraph()
    doc.add_heading("Out of Scope", level=2)
    future_phases = data.get("future_phases", [])
    if future_phases:
        for phase in future_phases:
            doc.add_paragraph(phase.get("name", "Future Phase"), style="List Bullet")
            for task in phase.get("tasks", []):
                doc.add_paragraph(f"  • {task.get('name', '')}", style="List Bullet 2")
    else:
        doc.add_paragraph("(None specified)")
    doc.add_paragraph()


def add_estimate_tables(doc, data: dict):
    doc.add_heading("3. Estimate Breakdown", level=1)
    active_roles, role_names = get_roles(data)
    headers = (
        ["ID", "Task"] + [role_names.get(r, r) for r in active_roles] + ["Total MD", "SP", "Notes"]
    )
    for option in data.get("options", []):
        subtitle = f" ({option['subtitle']})" if option.get("subtitle") else ""
        doc.add_heading(
            f"Option {option.get('id', 'A')}: {option.get('name', 'Option')}{subtitle}", level=2
        )
        tasks = flatten_tasks(option)
        if not tasks:
            doc.add_paragraph("(No tasks)")
            continue
        rows = []
        for task in tasks:
            mds = [task_role_md(task, r) or 0 for r in active_roles]
            row = [task.get("id", ""), task.get("name", "")]
            row += [f"{md:.1f}" if md else "" for md in mds]
            row += [
                f"{task.get('total_md', sum(mds)):.1f}",
                str(task.get("story_points", "") or ""),
                task.get("notes", ""),
            ]
            rows.append(row)
        add_styled_table(doc, headers, rows)
        if option.get("ai_reduction_pct"):
            note = doc.add_paragraph()
            note.add_run(
                f"Note: AI-assisted development with {option['ai_reduction_pct']}% reduction applied."
            )
            note.runs[0].font.italic = True
        doc.add_paragraph()


def add_option_comparison(doc, data: dict):
    options = data.get("options", [])
    if len(options) < 2:
        return
    doc.add_heading("4. Option Comparison", level=1)
    headers = ["Option", "Total MD", "Scope Notes"]
    rows = []
    for option in options:
        option_id = option.get("id", "")
        option_name = option.get("name", "")
        total_md = sum_option_md(option)
        scope_notes = []
        if option.get("scope_reductions"):
            scope_notes = [r.get("item", "") for r in option["scope_reductions"]]
        scope_text = "; ".join(scope_notes) if scope_notes else "Full scope"
        rows.append([f"{option_id}: {option_name}", f"{total_md:.1f}", scope_text])
    add_styled_table(doc, headers, rows)
    doc.add_paragraph()


def add_team_composition(doc, data: dict):
    doc.add_heading("5. Team Composition", level=1)
    active_roles, role_names = get_roles(data)
    if not active_roles:
        doc.add_paragraph("(No roles specified)")
        return
    headers = ["Role", "Name"]
    rows = [[role_names.get(r, r), role_names.get(r, r)] for r in active_roles]
    add_styled_table(doc, headers, rows)
    doc.add_paragraph()


def add_risks(doc, data: dict):
    doc.add_heading("6. Risk Assessment", level=1)
    risks = data.get("risks", [])
    if not risks:
        doc.add_paragraph("(No risks identified)")
        return
    _keys = ["id", "description", "category", "impact", "mitigation"]
    rows = [[str(risk.get(k, "")) for k in _keys] for risk in risks]
    add_styled_table(doc, ["ID", "Description", "Category", "Impact", "Mitigation"], rows)
    doc.add_paragraph()


def _add_numbered_list(doc, heading: str, items: list, level: int = 2):
    doc.add_heading(heading, level=level)
    if items:
        for i, item in enumerate(items, 1):
            doc.add_paragraph(f"{i}. {item}")
    else:
        doc.add_paragraph("(None)")


def add_assumptions(doc, data: dict):
    doc.add_heading("7. Assumptions & Constraints", level=1)
    _add_numbered_list(doc, "Assumptions", data.get("assumptions", []))
    doc.add_paragraph()
    doc.add_heading("TBD Items", level=2)
    tbd_items = data.get("tbd_items", [])
    if tbd_items:
        _keys = ["item", "status", "risk_impact", "recommendation"]
        rows = [[item.get(k, "") for k in _keys] for item in tbd_items]
        add_styled_table(doc, ["Item", "Status", "Risk Impact", "Recommendation"], rows)
    else:
        doc.add_paragraph("(None)")
    doc.add_paragraph()
    _add_numbered_list(doc, "Unresolved Questions", data.get("unresolved_questions", []))
