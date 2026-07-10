"""Assumptions and preconditions JSON to DOCX generator."""

from datetime import datetime
from pathlib import Path


def render(data: dict, output_path, config: dict | None = None) -> None:
    """Render assumptions from estimate JSON to DOCX.

    Args:
        data: Estimate JSON data containing assumptions, unresolved_questions, and tbd_items
        output_path: Path to write DOCX file
        config: Optional configuration (not used)
    """
    from docx import Document

    doc = Document()

    project_name = data.get("project_name", "Untitled Project")
    doc.add_heading(f"{project_name} — Assumptions & Preconditions", level=0)

    _add_metadata_table(doc, data)

    assumptions = data.get("assumptions", [])
    if assumptions:
        _add_assumptions_section(doc, assumptions)

    unresolved_questions = data.get("unresolved_questions", [])
    if unresolved_questions:
        _add_unresolved_questions_section(doc, unresolved_questions)

    tbd_items = data.get("tbd_items", [])
    if tbd_items:
        _add_tbd_items_section(doc, tbd_items)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _add_metadata_table(doc, data: dict) -> None:
    """Add metadata table with project info."""
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"

    rows_data = [
        ("Project", data.get("project_name", "N/A")),
        ("Date", data.get("generated_date", datetime.now().strftime("%Y-%m-%d"))),
        ("Version", data.get("version", "1.0")),
        ("Estimator", data.get("estimator", "N/A")),
    ]

    for row_idx, (label, value) in enumerate(rows_data):
        table.rows[row_idx].cells[0].text = label
        table.rows[row_idx].cells[1].text = str(value)
        table.rows[row_idx].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()


def _add_assumptions_section(doc, assumptions) -> None:
    """Add assumptions section with numbered list or table."""
    doc.add_heading("Assumptions", level=1)

    if not assumptions:
        return

    if isinstance(assumptions[0], str):
        for assumption in assumptions:
            doc.add_paragraph(assumption, style="List Number")
    else:
        table = doc.add_table(rows=len(assumptions) + 1, cols=3)
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Assumption"
        header_cells[1].text = "Category"
        header_cells[2].text = "Impact"

        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        for idx, assumption in enumerate(assumptions, start=1):
            row = table.rows[idx]
            row.cells[0].text = assumption.get("text", str(assumption))
            row.cells[1].text = assumption.get("category", "")
            row.cells[2].text = assumption.get("impact", "")

    doc.add_paragraph()


def _add_unresolved_questions_section(doc, questions: list) -> None:
    """Add unresolved questions section."""
    doc.add_heading("Unresolved Questions", level=1)

    for question in questions:
        doc.add_paragraph(question, style="List Number")

    doc.add_paragraph()


def _add_tbd_items_section(doc, tbd_items: list) -> None:
    """Add TBD items section as table."""
    doc.add_heading("TBD Items", level=1)

    if not tbd_items:
        return

    table = doc.add_table(rows=len(tbd_items) + 1, cols=4)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Status"
    header_cells[2].text = "Risk Impact"
    header_cells[3].text = "Recommendation"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, item in enumerate(tbd_items, start=1):
        row = table.rows[idx]
        row.cells[0].text = item.get("item", "")
        row.cells[1].text = item.get("status", "")
        row.cells[2].text = item.get("risk_impact", "")
        row.cells[3].text = item.get("recommendation", "")

    doc.add_paragraph()
